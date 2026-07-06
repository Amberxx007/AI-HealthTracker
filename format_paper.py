
import os
from docx import Document
from docx.styles.styles import Styles
from pathlib import Path

def copy_style(source_style, target_doc):
    """Copies a style from a source style object to a target document."""
    # The check for _Style is removed as it's an internal class.
    # We rely on duck typing; if it has .name and .type, it's a style.
    if not hasattr(source_style, 'name') or not hasattr(source_style, 'type'):
        return
    
    if source_style.name not in target_doc.styles:
        target_style = target_doc.styles.add_style(source_style.name, source_style.type)
        
        # Copy font properties
        if source_style.font:
            target_style.font.name = source_style.font.name
            target_style.font.size = source_style.font.size
            target_style.font.bold = source_style.font.bold
            target_style.font.italic = source_style.font.italic
            target_style.font.underline = source_style.font.underline
            target_style.font.color.rgb = source_style.font.color.rgb

        # Copy paragraph format properties
        if source_style.paragraph_format:
            target_style.paragraph_format.alignment = source_style.paragraph_format.alignment
            target_style.paragraph_format.first_line_indent = source_style.paragraph_format.first_line_indent
            target_style.paragraph_format.left_indent = source_style.paragraph_format.left_indent
            target_style.paragraph_format.right_indent = source_style.paragraph_format.right_indent
            target_style.paragraph_format.space_before = source_style.paragraph_format.space_before
            target_style.paragraph_format.space_after = source_style.paragraph_format.space_after
            target_style.paragraph_format.line_spacing = source_style.paragraph_format.line_spacing

def main():
    """
    Main function to format the research paper according to the B.Tech project template.
    """
    workspace_root = Path(r"d:\ai-doctor-v3")
    template_path = workspace_root / "Project Format for B.Tech 6th Sem.docx"
    source_path = workspace_root / "research_paper_final_clean (1).docx"
    output_path = workspace_root / "BTech_Final_Project_Report.docx"

    print("--- Starting Document Formatting ---")

    # --- Verification Step ---
    if not template_path.exists():
        print(f"ERROR: Template file not found at: {template_path}")
        return
    if not source_path.exists():
        print(f"ERROR: Source content file not found at: {source_path}")
        return

    print(f"Template found: {template_path.name}")
    print(f"Source content found: {source_path.name}")

    # --- Load Documents ---
    try:
        template_doc = Document(template_path)
        source_doc = Document(source_path)
        print("Successfully loaded both template and source documents.")
    except Exception as e:
        print(f"An error occurred while loading documents: {e}")
        return

    # --- Create New Document from Template ---
    # We will build the new document by iterating through the template and replacing/adding content.
    # This is more robust for preserving the initial pages' strict formatting.
    new_doc = Document()

    # Copy all styles from template to new document
    for style in template_doc.styles:
        copy_style(style, new_doc)

    # --- Content Transfer and Formatting ---
    # This is a simplified transfer. A full, perfect transfer is extremely complex.
    # We will iterate through the source and add content to the new doc, applying styles.
    
    # For this demonstration, let's copy the first few pages of the template exactly
    # to preserve the strict formatting, then append the content from the source.
    
    # A placeholder for where content from the source doc should start.
    # In a real scenario, you'd map sections. Let's find the "Introduction".
    template_content_start_index = 0
    for i, para in enumerate(template_doc.paragraphs):
        if "introduction" in para.text.lower():
            template_content_start_index = i
            break
    
    print(f"Template content will be used up to paragraph {template_content_start_index}.")

    # 1. Copy initial pages from template
    for i in range(template_content_start_index):
        para = template_doc.paragraphs[i]
        new_para = new_doc.add_paragraph(style=para.style.name if para.style.name in new_doc.styles else None)
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            # Copy run formatting
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size

    print("Copied initial pages from template.")

    # 2. Append content from source document
    print("Appending content from source document...")
    for element in source_doc.element.body:
        if element.tag.endswith('p'):
            source_para = [p for p in source_doc.paragraphs if p._p == element][0]
            style_name = source_para.style.name
            
            # Try to use the same style, otherwise use default
            target_style = new_doc.styles[style_name] if style_name in new_doc.styles else None
            new_para = new_doc.add_paragraph(style=target_style)
            
            for run in source_para.runs:
                new_run = new_para.add_run(run.text)
                # Basic formatting copy
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
        
        elif element.tag.endswith('tbl'):
            # Table handling is more complex, this is a basic copy
            source_table = [t for t in source_doc.tables if t._tbl == element][0]
            new_table = new_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns), style='Table Grid')
            for i, row in enumerate(source_table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.cell(i, j).text = cell.text

    print("Finished appending source content.")

    # --- Save Final Document ---
    try:
        new_doc.save(output_path)
        print(f"SUCCESS: Final project report saved to: {output_path}")
    except Exception as e:
        print(f"An error occurred while saving the final document: {e}")

if __name__ == "__main__":
    main()
